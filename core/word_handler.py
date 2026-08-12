"""
Módulo para criação e manipulação de documentos Word no projeto MARIA.
Usa a biblioteca python-docx para gerar arquivos .docx.
"""

import os
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.file_utils import garantir_pasta_arquivos, gerar_nome_unico
from core.config import PASTA_ARQUIVOS_GERADOS

logger = logging.getLogger(__name__)


def criar_documento_real(nome_arquivo: str, titulo: str, conteudo: str) -> str:
    """
    Cria um documento Word com título e conteúdo textual.
    
    Args:
        nome_arquivo: Nome do arquivo (sem extensão)
        titulo: Título principal do documento
        conteudo: Texto completo, com parágrafos separados por duas quebras de linha
        
    Returns:
        Caminho absoluto do arquivo criado
        
    Raises:
        PermissionError: Se não houver permissão de escrita
        OSError: Se houver erro de disco
    """
    try:
        # Gerar nome único para evitar sobrescrita
        nome_final = gerar_nome_unico(nome_arquivo, ".docx")
        pasta_absoluta = garantir_pasta_arquivos()
        caminho_completo = os.path.join(pasta_absoluta, nome_final)
        
        # Criar novo documento
        doc = Document()
        
        # Adicionar título principal
        titulo_para = doc.add_heading(titulo, level=1)
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if conteudo and conteudo.strip():
            for paragrafo in conteudo.split("\n\n"):
                if paragrafo.strip():
                    doc.add_paragraph(paragrafo.strip())
        else:
            doc.add_paragraph(
                "Este é um documento gerado pela MARIA. "
                "Você pode editar este conteúdo conforme necessário."
            )
        
        # Salvar documento
        doc.save(caminho_completo)
        
        logger.info(f"Documento criado: {caminho_completo}")
        return caminho_completo
        
    except PermissionError as e:
        logger.error(f"Permissão negada ao criar documento: {e}")
        raise PermissionError(
            f"Não foi possível salvar o arquivo. Verifique as permissões da pasta '{PASTA_ARQUIVOS_GERADOS}'."
        ) from e
        
    except OSError as e:
        logger.error(f"Erro de disco ao criar documento: {e}")
        raise OSError(
            "Não foi possível salvar o arquivo. Verifique se há espaço em disco disponível."
        ) from e
        
    except Exception as e:
        logger.error(f"Erro inesperado ao criar documento: {e}")
        raise
